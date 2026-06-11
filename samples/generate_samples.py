"""Generate sample DOCX template files with 【变量名】 placeholders."""
from docx import Document
from docx.shared import Pt
import os

SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "samples")


def create_shareholder_resolution():
    doc = Document()
    doc.add_heading("股东会决议签字页", level=1)
    doc.add_paragraph(
        f"【公司名称】于【年份】年【月份】月【日期】日召开股东会，"
        f"出席股东：【股东姓名】。"
    )
    doc.add_paragraph(
        f"经全体股东一致同意，决议如下："
    )
    doc.add_paragraph(
        f"股东签字：【股东姓名】"
    )
    doc.add_paragraph(
        f"日期：【年份】年【月份】月【日期】日"
    )
    path = os.path.join(SAMPLES_DIR, "签字页模板_股东会决议.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_board_resolution():
    doc = Document()
    doc.add_heading("董事会决议签字页", level=1)
    doc.add_paragraph(
        f"【公司名称】于【年份】年【月份】月【日期】日召开董事会，"
        f"出席董事：【董事姓名】。"
    )
    doc.add_paragraph(
        f"经全体董事一致同意，决议如下："
    )
    doc.add_paragraph(
        f"董事签字：【董事姓名】"
    )
    doc.add_paragraph(
        f"日期：【年份】年【月份】月【日期】日"
    )
    path = os.path.join(SAMPLES_DIR, "签字页模板_董事会决议.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_lawyer_witness_letter():
    doc = Document()
    doc.add_heading("律师见证函签字页", level=1)
    doc.add_paragraph(
        f"【律所名称】受【公司名称】委托，指派【律师姓名】律师"
        f"对本次会议进行见证。"
    )
    doc.add_paragraph(
        f"见证律师签字：【律师姓名】"
    )
    doc.add_paragraph(
        f"律所盖章：【律所名称】"
    )
    doc.add_paragraph(
        f"日期：【年份】年【月份】月【日期】日"
    )
    path = os.path.join(SAMPLES_DIR, "签字页模板_律师见证函.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    os.makedirs(SAMPLES_DIR, exist_ok=True)
    create_shareholder_resolution()
    create_board_resolution()
    create_lawyer_witness_letter()
    print("All sample templates generated.")