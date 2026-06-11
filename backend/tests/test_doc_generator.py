import os

from docx import Document

from app.utils.doc_generator import batch_generate_docx, generate_docx, preview_docx


def _create_test_docx(variables_text: str, tmp_dir: str) -> str:
    doc = Document()
    doc.add_paragraph(variables_text)
    path = os.path.join(tmp_dir, "test_template.docx")
    doc.save(path)
    return path


class TestGenerateDocx:
    def test_simple_replace(self, tmp_path):
        template_path = _create_test_docx("甲方：【公司名称】", str(tmp_path))
        output_path = os.path.join(str(tmp_path), "output.docx")
        result_path = generate_docx(template_path, {"公司名称": "测试公司"}, output_path)
        assert os.path.exists(result_path)
        doc = Document(result_path)
        assert "测试公司" in doc.paragraphs[0].text
        assert "【公司名称】" not in doc.paragraphs[0].text

    def test_multiple_replace(self, tmp_path):
        template_path = _create_test_docx("【公司名称】【法定代表人】【日期】", str(tmp_path))
        output_path = os.path.join(str(tmp_path), "output.docx")
        generate_docx(
            template_path,
            {"公司名称": "XX科技", "法定代表人": "张三", "日期": "2024-01-01"},
            output_path,
        )
        doc = Document(output_path)
        text = doc.paragraphs[0].text
        assert "XX科技" in text
        assert "张三" in text
        assert "2024-01-01" in text

    def test_unfilled_variable_preserved(self, tmp_path):
        template_path = _create_test_docx("【公司名称】【未填变量】", str(tmp_path))
        output_path = os.path.join(str(tmp_path), "output.docx")
        generate_docx(template_path, {"公司名称": "XX科技"}, output_path)
        doc = Document(output_path)
        text = doc.paragraphs[0].text
        assert "XX科技" in text
        assert "【未填变量】" in text

    def test_output_file_valid_docx(self, tmp_path):
        template_path = _create_test_docx("【变量】", str(tmp_path))
        result_path = generate_docx(template_path, {"变量": "值"})
        assert os.path.exists(result_path)
        assert result_path.endswith(".docx")
        doc = Document(result_path)
        assert len(doc.paragraphs) > 0


class TestBatchGenerateDocx:
    def test_batch(self, tmp_path):
        template_path = _create_test_docx("【名称】", str(tmp_path))
        output_dir = os.path.join(str(tmp_path), "batch_output")
        variables_list = [
            {"名称": "第一份"},
            {"名称": "第二份"},
            {"名称": "第三份"},
        ]
        paths = batch_generate_docx(template_path, variables_list, output_dir)
        assert len(paths) == 3
        for p in paths:
            assert os.path.exists(p)


class TestPreviewDocx:
    def test_preview_with_replacement(self, tmp_path):
        template_path = _create_test_docx("甲方：【公司名称】同意", str(tmp_path))
        result = preview_docx(template_path, {"公司名称": "XX科技"})
        assert "XX科技" in result
        assert "【公司名称】" not in result

    def test_preview_without_replacement(self, tmp_path):
        template_path = _create_test_docx("甲方：【公司名称】同意", str(tmp_path))
        result = preview_docx(template_path, {})
        assert "【公司名称】" in result
