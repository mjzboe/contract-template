"""文档生成器：读取 DOCX 模板，替换【变量名】为实际值，生成新文档"""

import re
import tempfile

from docx import Document

from docx.oxml.ns import qn

from app.config import resolve_file_path
from app.utils.variable_parser import CHINESE_BRACKET_PATTERN


def generate_docx(
    template_path: str,
    variables: dict[str, str],
    output_path: str | None = None,
) -> str:
    """根据模板和变量值生成新文档

    Args:
        template_path: 模板文件路径
        variables: 变量名 → 值 的映射，如 {"公司名称": "XX科技有限公司"}
        output_path: 输出路径，None 则用临时文件

    Returns:
        生成文件的路径
    """
    resolved = resolve_file_path(template_path)
    doc = Document(resolved)

    # 确保所有 run 都有东亚字体设置，避免 LibreOffice 转换 PDF 时中文变方框
    _ensure_cjk_font(doc)

    # 替换段落中的变量
    for paragraph in doc.paragraphs:
        _replace_paragraph(paragraph, variables)

    # 替换表格中的变量
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_paragraph(paragraph, variables)

    # 替换页眉页脚中的变量
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _replace_paragraph(paragraph, variables)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _replace_paragraph(paragraph, variables)

    # 确定输出路径
    if not output_path:
        fd, output_path = tempfile.mkstemp(suffix=".docx")
        import os
        os.close(fd)

    doc.save(output_path)
    return output_path


def batch_generate_docx(
    template_path: str,
    variables_list: list[dict[str, str]],
    output_dir: str,
) -> list[str]:
    """批量生成文档（同一模板，多组变量）

    Returns:
        生成的文件路径列表
    """
    import os
    os.makedirs(output_dir, exist_ok=True)

    output_paths = []
    for i, variables in enumerate(variables_list, 1):
        filename = f"generated_{i}.docx"
        output_path = os.path.join(output_dir, filename)
        generate_docx(template_path, variables, output_path)
        output_paths.append(output_path)

    return output_paths


def preview_docx(template_path: str, variables: dict[str, str]) -> str:
    """预览：替换变量后返回文档的纯文本内容"""
    resolved = resolve_file_path(template_path)
    doc = Document(resolved)

    # 替换段落
    for paragraph in doc.paragraphs:
        _replace_paragraph(paragraph, variables)

    # 提取纯文本
    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    return "\n".join(lines)


def _replace_paragraph(paragraph, variables: dict[str, str]):
    """替换段落中所有 run 的【变量名】为实际值

    注意：Word 的 run 可能会把一个【变量名】拆分成多个 run，
    所以先合并段落文本，再整体替换。
    """
    # 检查段落是否包含变量标记
    full_text = paragraph.text
    if "【" not in full_text and "】" not in full_text:
        return

    # 执行替换
    new_text = _replace_variables_in_text(full_text, variables)
    if new_text == full_text:
        return

    # 清除原有 runs，写入新文本
    # 保留第一个 run 的格式
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def _replace_variables_in_text(text: str, variables: dict[str, str]) -> str:
    """替换文本中的【变量名】"""

    def replacer(match):
        var_name = match.group(1)
        return variables.get(var_name, match.group(0))  # 未找到则保留原样

    return CHINESE_BRACKET_PATTERN.sub(replacer, text)


CJK_FONT = "Noto Sans CJK SC"


def _ensure_cjk_font(doc: Document):
    """为文档中所有 run 设置东亚字体，避免 LibreOffice 转换 PDF 时中文变方框。

    Word 默认用宋体/等线等中文字体，但 Docker 容器中没有这些字体，
    LibreOffice 会 fallback 到不支持中文的字体。显式设置为容器中已安装的
    Noto Sans CJK SC 可解决此问题。
    """
    for paragraph in doc.paragraphs:
        _set_run_cjk_font(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_run_cjk_font(paragraph)
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer:
                for paragraph in header_footer.paragraphs:
                    _set_run_cjk_font(paragraph)


def _set_run_cjk_font(paragraph):
    """为段落中每个 run 设置东亚字体"""
    for run in paragraph.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = run._element.makeelement(qn("w:rPr"), {})
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        if not rfonts.get(qn("w:eastAsia")):
            rfonts.set(qn("w:eastAsia"), CJK_FONT)
