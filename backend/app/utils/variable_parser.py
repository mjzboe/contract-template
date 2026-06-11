"""变量解析器：从 DOCX 模板中提取【变量名】格式的变量"""

import re
from dataclasses import dataclass, field

# 中文方括号变量（项目实际格式）
CHINESE_BRACKET_PATTERN = re.compile(r"【(.+?)】")

# 花括号变量（PRD 描述格式，兼容支持）
CURLY_BRACE_PATTERN = re.compile(r"\{\{(\w+)(?:\|([^}]+))?(?::(\w+)(?::([^}]+))?)?\}\}")


@dataclass
class VariableInfo:
    """解析出的变量信息"""

    name: str
    display_name: str = ""
    var_type: str = "text"
    default_value: str = ""
    validation_rule: str = ""
    occurrences: int = 1

    def __post_init__(self):
        if not self.display_name:
            self.display_name = self.name


def extract_variables_from_text(text: str) -> list[VariableInfo]:
    """从文本中提取所有变量，去重并统计出现次数"""
    var_map: dict[str, VariableInfo] = {}

    # 1. 提取中文方括号变量 【变量名】
    for match in CHINESE_BRACKET_PATTERN.finditer(text):
        name = match.group(1).strip()
        if name and name not in var_map:
            var_map[name] = VariableInfo(name=name)
        elif name:
            var_map[name].occurrences += 1

    # 2. 提取花括号变量 {{变量名}} / {{变量名|默认值}} / {{变量名:type:rule}}
    for match in CURLY_BRACE_PATTERN.finditer(text):
        name = match.group(1)
        default = match.group(2) or ""
        var_type = match.group(3) or "text"
        rule = match.group(4) or ""

        if name in var_map:
            # 已存在则更新类型信息（花括号格式更丰富）
            var_map[name].var_type = var_type
            var_map[name].default_value = default
            var_map[name].validation_rule = rule
            var_map[name].occurrences += 1
        else:
            var_map[name] = VariableInfo(
                name=name,
                var_type=var_type,
                default_value=default,
                validation_rule=rule,
            )

    return list(var_map.values())


def extract_variables_from_docx(file_path: str) -> list[VariableInfo]:
    """从 DOCX 文件中提取所有变量"""
    from docx import Document

    doc = Document(file_path)
    all_text_parts: list[str] = []

    # 提取段落文本
    for paragraph in doc.paragraphs:
        all_text_parts.append(paragraph.text)

    # 提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text_parts.append(cell.text)

    # 提取页眉页脚
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                all_text_parts.append(p.text)
        if section.footer:
            for p in section.footer.paragraphs:
                all_text_parts.append(p.text)

    full_text = "\n".join(all_text_parts)
    return extract_variables_from_text(full_text)


def deduplicate_variables(variables_list: list[list[VariableInfo]]) -> list[VariableInfo]:
    """跨模板变量去重：多个模板的变量合并，同名变量只保留一个，occurrences 累加"""
    merged: dict[str, VariableInfo] = {}

    for variables in variables_list:
        for var in variables:
            if var.name in merged:
                merged[var.name].occurrences += var.occurrences
                # 保留更丰富的类型信息
                if var.var_type != "text" and merged[var.name].var_type == "text":
                    merged[var.name].var_type = var.var_type
                    merged[var.name].default_value = var.default_value
                    merged[var.name].validation_rule = var.validation_rule
            else:
                merged[var.name] = VariableInfo(
                    name=var.name,
                    display_name=var.display_name,
                    var_type=var.var_type,
                    default_value=var.default_value,
                    validation_rule=var.validation_rule,
                    occurrences=var.occurrences,
                )

    return list(merged.values())
